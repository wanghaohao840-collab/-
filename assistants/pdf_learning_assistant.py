# assistants/pdf_learning_assistant.py

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from threading import RLock
from typing import Optional, Dict, Any, Sequence

from app.history import HistoryRepository
from app.summary_tasks import SummaryTaskManager
from hello_agents.memory.base import MemoryConfig
from hello_agents.memory.rag.result_utils import (
    MAX_SELECTED_DOCUMENTS,
    resolve_qa_mode,
)
from hello_agents.tools.builtin.memory_tool import MemoryTool
from hello_agents.tools.builtin.rag_tool import RAGTool
from assistants.document_selection import build_document_scope


class ImportRAGError(RuntimeError):
    """Structured RAG import failure consumed by the background runner."""

    def __init__(self, action_result: Any):
        message = str(
            getattr(action_result, "error", "")
            or getattr(action_result, "message", "")
            or "RAG document import failed"
        )
        super().__init__(message)
        self.error_code = str(
            getattr(action_result, "error_code", "") or "rag_operation"
        )
        self.retryable = bool(getattr(action_result, "retryable", False))
        self.action_result = action_result


class ImportMemoryEventError(RuntimeError):
    """Import completed durably but its idempotent memory event failed."""

    error_code = "memory_import_event"
    retryable = True

    def __init__(self):
        super().__init__("Failed to record the import memory event")


class PDFLearningAssistant:
    """PDF 学习助手

    功能：
    1. 导入 PDF 文档
    2. 基于 PDF 内容问答
    3. 记录用户问题
    4. 添加学习笔记
    5. 回忆历史学习内容
    6. 查看学习统计
    """

    def __init__(
        self,
        user_id: str = "user123",
        runtime_dir: Optional[str | Path] = None,
        runtime: Any = None,
    ):
        self.user_id = user_id
        self.session_id = datetime.now().strftime("session_%Y%m%d_%H%M%S")
        self.runtime = runtime
        self.runtime_dir = Path(runtime_dir or getattr(getattr(runtime, "paths", None), "root", None) or "./.runtime").resolve()
        self.memory_dir = self.runtime_dir / "memory"
        self.rag_dir = self.runtime_dir / "rag"
        self.report_dir = self.runtime_dir / "reports"
        self._lock = getattr(runtime, "lock", RLock())

        self.memory_dir.mkdir(parents=True, exist_ok=True)
        self.rag_dir.mkdir(parents=True, exist_ok=True)
        self.report_dir.mkdir(parents=True, exist_ok=True)

        if runtime is not None:
            self.memory_tool = runtime.memory_tool
        else:
            self.memory_tool = MemoryTool(
                user_id=user_id,
                memory_config=MemoryConfig(
                    database_path=str(self.memory_dir / f"memory_{user_id}.db")
                ),
                memory_types=["working", "episodic", "semantic"]
            )

        if runtime is not None:
            self.rag_tool = runtime.rag_tool
        else:
            self.rag_tool = RAGTool(
                knowledge_base_path=str(self.rag_dir),
                collection_name="pdf_learning_collection",
                rag_namespace=f"pdf_{user_id}",
                cache_path=str(self.rag_dir / "rag_cache.json"),
            )

        self.current_document: Optional[str] = None
        self.current_document_id: Optional[str] = None

        self.history_repository = getattr(runtime, "history", None)
        self.report_service = getattr(runtime, "reports", None)
        self.coordinator = getattr(runtime, "coordinator", None)
        self.history_path = Path(getattr(self.history_repository, "path", self._resolve_history_path(user_id)))
        if self.history_repository is None:
            self.history_repository = HistoryRepository(self.history_path)
        self.history_path.parent.mkdir(parents=True, exist_ok=True)

        self.history = self._load_history()

        self.stats = {
            "session_start": datetime.now().isoformat(),
            "documents_loaded": 0,
            "questions_asked": 0,
            "notes_added": 0,
        }
        self._summary_task_manager = SummaryTaskManager(max_workers=2)

    def close(self) -> None:
        """Release runtime resources such as SQLite connections."""

        manager = getattr(self, "_summary_task_manager", None)
        if manager is not None:
            manager.close()
        if self.runtime is not None:
            return

        rag_close = getattr(getattr(self, "rag_tool", None), "close", None)
        if callable(rag_close):
            rag_close()

        manager = getattr(getattr(self, "memory_tool", None), "memory_manager", None)
        memory_types = getattr(manager, "memory_types", {}) or {}
        for memory in memory_types.values():
            doc_store = getattr(memory, "doc_store", None)
            if doc_store is not None and hasattr(doc_store, "close"):
                doc_store.close()

    def __del__(self):
        try:
            self.close()
        except Exception:
            pass

    def _resolve_history_path(self, user_id: str) -> Path:
        if self.runtime_dir.name == str(user_id):
            return self.runtime_dir / "history.json"
        return self.memory_dir / f"learning_history_{user_id}.json"

    @property
    def _write_lock(self):
        # Some lightweight tests construct the assistant without __init__.
        if not hasattr(self, "_lock"):
            self._lock = RLock()
        return self._lock

    def load_document(
        self,
        pdf_path: str,
        document_id: Optional[str] = None,
        original_name: Optional[str] = None,
        import_task_id: Optional[str] = None,
        progress_callback: Any = None,
    ) -> str:
        """导入文档，支持 PDF / TXT / Markdown"""

        path = Path(pdf_path)

        if not path.exists():
            return f"❌ 文件不存在: {pdf_path}"

        suffix = path.suffix.lower()

        supported_suffixes = [".pdf", ".txt", ".md", ".markdown", ".docx"]

        if suffix not in supported_suffixes:
            return f"❌ 当前只支持 PDF / TXT / Markdown / Word 文件: {suffix}"

        document_id = document_id or path.stem
        document_name = original_name or path.name

        add_kwargs = {
            "file_path": str(path),
            "document_id": document_id,
            "metadata": {
                "user_id": self.user_id,
                "session_id": self.session_id,
                "document_id": document_id,
                "document_name": document_name,
                "document_path": str(path),
                "file_suffix": suffix,
                "source_type": "document",
                "import_task_id": import_task_id,
            },
        }
        if progress_callback is not None:
            add_kwargs["progress_callback"] = progress_callback

        history_item = {
            "document_id": document_id,
            "document_name": document_name,
            "document_path": str(path),
            "file_suffix": suffix,
            "session_id": self.session_id,
            "loaded_at": datetime.now().isoformat(),
            "import_task_id": import_task_id,
        }

        # ── critical section: RAG + History + import memory event ───
        with self._write_lock:
            latest_history = (
                self.history_repository.load() if import_task_id else self.history
            )
            existing = None
            if import_task_id:
                existing = next(
                    (
                        item
                        for item in latest_history["documents"]
                        if str(item.get("document_id", "")) == document_id
                    ),
                    None,
                )
            rag_existed_before: Optional[bool] = None
            try:
                pipeline = self.rag_tool._get_pipeline()
                list_document_ids = getattr(pipeline, "list_document_ids", None)
                if callable(list_document_ids):
                    rag_existed_before = document_id in list_document_ids()
            except Exception:
                pass
            rag_has_document = rag_existed_before is True

            already_imported = bool(
                import_task_id
                and existing
                and existing.get("import_task_id") == import_task_id
                and rag_has_document
            )
            rag_mutated = False
            rag_created_by_invocation = False
            if already_imported:
                result = f"✅ 文档已导入\n- document_id: {document_id}"
                committed_history = latest_history
            else:
                if hasattr(self.rag_tool, "execute_result"):
                    action_result = self.rag_tool.execute_result(
                        "add_document", **add_kwargs
                    )
                    result = action_result.message
                    if not action_result.success:
                        raise ImportRAGError(action_result)
                else:
                    result = self.rag_tool.execute("add_document", **add_kwargs)
                rag_mutated = True
                rag_created_by_invocation = rag_existed_before is False or (
                    import_task_id is None and rag_existed_before is None
                )

                # Commit History inside the same lock — fresh merge against
                # the latest persisted snapshot.
                try:
                    if self.coordinator is not None:
                        def upsert(data):
                            for index, item in enumerate(data["documents"]):
                                if str(item.get("document_id", "")) == document_id:
                                    data["documents"][index] = dict(history_item)
                                    return
                            data["documents"].append(dict(history_item))

                        committed_history = self.coordinator.update_history(upsert)
                    else:
                        committed_history = self.history_repository.upsert_document(
                            history_item
                        )
                except Exception:
                    # Only undo a RAG document written by this invocation.
                    if rag_mutated and rag_created_by_invocation:
                        if self.coordinator is not None:
                            self.coordinator.compensate_rag_add(
                                self.rag_tool,
                                document_id,
                                reason="history update failed",
                            )
                        else:
                            try:
                                self.rag_tool.execute(
                                    "delete_document", document_id=document_id
                                )
                            except Exception:
                                pass
                    raise

            memory_metadata = {
                "document_id": document_id,
                "document_name": document_name,
                "document_path": str(path),
                "file_suffix": suffix,
                "import_task_id": import_task_id,
            }
            try:
                if import_task_id and hasattr(
                    self.memory_tool, "ensure_import_event"
                ):
                    self.memory_tool.ensure_import_event(
                        import_task_id=import_task_id,
                        content=f"用户导入了文档：{document_name}",
                        metadata=memory_metadata,
                        session_id=self.session_id,
                    )
                else:
                    memory_result = self.memory_tool.execute(
                        "add",
                        content=f"用户导入了文档：{path.name}",
                        memory_type="episodic",
                        importance=0.8,
                        event_type="document_loaded",
                        session_id=self.session_id,
                        metadata=memory_metadata,
                    )
                    if isinstance(memory_result, str) and memory_result.startswith("❌"):
                        raise RuntimeError(memory_result)
            except ImportMemoryEventError:
                raise
            except Exception as error:
                raise ImportMemoryEventError() from error

        # ── outside lock: session-local state ────────────────────────
        self.history = committed_history
        self.current_document = str(path)
        self.current_document_id = document_id
        self.stats["documents_loaded"] += 1

        return result

    def ask(
        self,
        question: str,
        limit: int = 5,
        selected_documents: Optional[Sequence[str] | str] = None,
        mode: str = "auto",
        progress_callback: Any = None,
        cancel_event: Any = None,
        structured_output: bool = False,
    ) -> str:
        """Answer questions against the current document or an explicit set."""

        if not question.strip():
            return "❌ 问题不能为空"

        try:
            scope = build_document_scope(selected_documents)
        except ValueError as error:
            return f"❌ 文档选择无效: {error}"
        explicit_scope = selected_documents is not None

        if explicit_scope and not scope.document_ids:
            return "❌ 请选择至少 1 篇文档"
        if explicit_scope and len(scope.document_ids) > MAX_SELECTED_DOCUMENTS:
            return f"❌ 最多选择 {MAX_SELECTED_DOCUMENTS} 篇文档"

        if not explicit_scope and not self.current_document_id:
            return "❌ 当前没有选择文档。请先上传文档或在下拉框中选择一个文档。"

        try:
            selected_mode = resolve_qa_mode(question, mode)
        except ValueError as error:
            return f"❌ {error}"

        selected_ids = (
            scope.document_ids
            if explicit_scope
            else [self.current_document_id]
        )
        if selected_mode == "compare" and len(selected_ids) < 2:
            return "❌ 对比分析至少需要选择 2 篇文档"

        self.stats["questions_asked"] += 1

        self.memory_tool.execute(
            "add",
            content=f"用户提出问题：{question}",
            memory_type="working",
            importance=0.6,
            session_id=self.session_id,
        )

        rag_kwargs = {
            "query": question,
            "limit": limit,
            "min_score": 0.12,
            "mode": selected_mode,
        }
        if progress_callback is not None:
            rag_kwargs["progress_callback"] = progress_callback
        if cancel_event is not None:
            rag_kwargs["cancel_event"] = cancel_event
        if structured_output:
            rag_kwargs["structured_output"] = True

        if explicit_scope:
            rag_kwargs["document_ids"] = scope.document_ids
        else:
            rag_kwargs["document_id"] = self.current_document_id

        answer = self.rag_tool.execute("ask", **rag_kwargs)
        is_cancelled = getattr(cancel_event, "is_set", None)
        if callable(is_cancelled) and is_cancelled():
            self.stats["questions_asked"] = max(
                0, self.stats["questions_asked"] - 1
            )
            return answer

        document_label = (
            "; ".join(scope.labels or [])
            if explicit_scope
            else self.current_document
        )
        history_item = {
            "question": question,
            "answer": answer,
            "document": document_label,
            "session_id": self.session_id,
            "asked_at": datetime.now().isoformat(),
        }

        if explicit_scope:
            history_item["document_ids"] = scope.document_ids
            history_item["document_names"] = scope.document_names
            history_item["mode"] = selected_mode

        if explicit_scope:
            document_ids = scope.document_ids
            document_names = scope.document_names
        else:
            latest = self._load_latest_history()
            document_ids = [self.current_document_id]
            document_names = [
                item.get("document_name", self.current_document_id)
                for item in latest["documents"]
                if item.get("document_id") == self.current_document_id
            ] or [Path(self.current_document or "").name]
        history_item["document_ids"] = document_ids
        history_item["document_names"] = document_names
        history_item["mode"] = selected_mode
        self._update_history(lambda history: history["questions"].append(history_item))

        self.memory_tool.execute(
            "add",
            content=f"用户针对文档提问：{question}\n系统回答：{answer[:300]}",
            memory_type="episodic",
            importance=0.7,
            event_type="pdf_qa",
            session_id=self.session_id,
        )

        return answer

    def start_summary_task(
        self,
        question: str,
        selected_documents: Sequence[str] | str,
        limit: int = 5,
    ) -> Dict[str, Any]:
        if not question or not question.strip():
            raise ValueError("question cannot be empty")
        scope = build_document_scope(selected_documents)
        if not scope.document_ids:
            raise ValueError("select at least one document")
        if len(scope.document_ids) > MAX_SELECTED_DOCUMENTS:
            raise ValueError(
                f"select at most {MAX_SELECTED_DOCUMENTS} documents"
            )
        manager = self._get_summary_task_manager()

        def run(progress_callback, cancel_event):
            return self.ask(
                question,
                limit=limit,
                selected_documents=list(selected_documents)
                if not isinstance(selected_documents, str)
                else selected_documents,
                mode="summary",
                progress_callback=progress_callback,
                cancel_event=cancel_event,
            )

        return manager.start(run, total=len(scope.document_ids))

    def get_summary_task(self, task_id: str) -> Dict[str, Any]:
        return self._get_summary_task_manager().get(task_id)

    def cancel_summary_task(self, task_id: str) -> Dict[str, Any]:
        return self._get_summary_task_manager().cancel(task_id)

    def _get_summary_task_manager(self) -> SummaryTaskManager:
        manager = getattr(self, "_summary_task_manager", None)
        if manager is None:
            manager = SummaryTaskManager(max_workers=2)
            self._summary_task_manager = manager
        return manager

    def search(
        self,
        query: str,
        limit: int = 5,
        selected_documents: Optional[Sequence[str] | str] = None,
    ) -> str:
        """Search the current document or an explicit document set."""

        if not query or not query.strip():
            return "❌ 检索关键词不能为空"

        try:
            scope = build_document_scope(selected_documents)
        except ValueError as error:
            return f"❌ 文档选择无效: {error}"
        explicit_scope = selected_documents is not None

        if explicit_scope and not scope.document_ids:
            return "❌ 请选择至少 1 篇文档"
        if explicit_scope and len(scope.document_ids) > MAX_SELECTED_DOCUMENTS:
            return f"❌ 最多选择 {MAX_SELECTED_DOCUMENTS} 篇文档"

        if not explicit_scope and not self.current_document_id:
            return "❌ 当前没有选择文档。请先上传文档或在下拉框中选择一个文档。"

        rag_kwargs = {
            "query": query,
            "limit": limit,
            "min_score": 0.08,
        }

        if explicit_scope:
            rag_kwargs["document_ids"] = scope.document_ids
        else:
            rag_kwargs["document_id"] = self.current_document_id

        return self.rag_tool.execute("search", **rag_kwargs)

    def add_note(self, note: str, concept: Optional[str] = None) -> str:
        """添加学习笔记"""

        if not note.strip():
            return "❌ 笔记内容不能为空"

        self.stats["notes_added"] += 1

        content = note
        if concept:
            content = f"关于【{concept}】的学习笔记：{note}"

        result = self.memory_tool.execute(
            "add",
            content=content,
            memory_type="semantic",
            importance=0.85,
            knowledge_type="learning_note",
            concept=concept or "",
            session_id=self.session_id
        )

        history_item = {
            "concept": concept or "",
            "note": note,
            "content": content,
            "session_id": self.session_id,
            "created_at": datetime.now().isoformat()
        }
        self._update_history(lambda history: history["notes"].append(history_item))

        return result

    def clear_all_notes(self) -> str:
        """清空全部学习笔记：只清理学习历史中的 notes，不影响 PDF 文档和问答历史"""

        latest = self._load_latest_history()
        removed_notes = len(latest.get("notes", []))

        # 1. 清空本地学习历史中的学习笔记
        self._update_history(lambda history: history.__setitem__("notes", []))

        # 2. 保存学习历史 JSON

        # 3. 重置统计
        if "notes_added" in self.stats:
            self.stats["notes_added"] = 0

        if "concepts_learned" in self.stats:
            self.stats["concepts_learned"] = 0

        return (
            "✅ 已清空全部学习笔记\n\n"
            f"- 删除历史学习笔记: {removed_notes} 条\n"
            "- PDF 文档和问答历史未删除\n"
            "- 当前 RAG 知识库未删除\n"
            "- 当前版本仅清理本地学习历史 notes"
        )

    def recall(self, query: str, limit: int = 5) -> str:
        """回忆历史学习内容：同时查询 MemoryTool 和本地学习历史 JSON"""

        if not query or not query.strip():
            return "❌ 回忆关键词不能为空"

        query = query.strip()

        # 1. 先查当前运行中的 MemoryTool
        memory_result = self.memory_tool.execute(
            "search",
            query=query,
            limit=limit
        )

        # 2. 再查本地 JSON 历史
        history_hits = []

        documents = self.history.get("documents", [])
        questions = self.history.get("questions", [])
        notes = self.history.get("notes", [])

        # 查历史文档
        for item in documents:
            text = f"{item.get('document_name', '')} {item.get('document_path', '')}"
            if query in text:
                history_hits.append(
                    f"[历史文档] {item.get('document_name')} | {item.get('loaded_at')}"
                )

        # 查历史问答
        for item in questions:
            question = str(item.get("question", ""))
            answer = str(item.get("answer", ""))

            if query in question or query in answer:
                short_answer = answer[:300].replace("\n", " ")
                history_hits.append(
                    f"[历史问答] 问题：{question}\n回答摘要：{short_answer}..."
                )

        # 查历史笔记
        for item in notes:
            concept = str(item.get("concept", ""))
            note = str(item.get("note", ""))
            content = str(item.get("content", ""))

            if query in concept or query in note or query in content:
                history_hits.append(
                    f"[历史笔记] 【{concept or '未命名概念'}】{note}"
                )

        # 3. 组合结果
        lines = []

        lines.append("一、当前记忆系统检索结果")
        lines.append(memory_result)

        lines.append("\n二、本地历史记录检索结果")

        if history_hits:
            for i, item in enumerate(history_hits[:limit], start=1):
                lines.append(f"{i}. {item}")
        else:
            lines.append(f"未在本地历史记录中找到与「{query}」相关的内容")

        return "\n".join(lines)

    def get_stats(self) -> str:
        """查看学习统计"""

        memory_summary = self.memory_tool.execute("summary")
        rag_stats = self.rag_tool.execute("stats")

        return (
            "📘 PDF 学习助手统计:\n"
            f"- user_id: {self.user_id}\n"
            f"- session_id: {self.session_id}\n"
            f"- 当前文档: {self.current_document or '暂无'}\n"
            f"- 当前文档ID: {self.current_document_id or '暂无'}\n"
            f"- 已导入文档数: {self.stats['documents_loaded']}\n"
            f"- 提问次数: {self.stats['questions_asked']}\n"
            f"- 学习笔记数: {self.stats['notes_added']}\n\n"
            f"{memory_summary}\n\n"
            f"{rag_stats}"
        )

    def generate_report(self) -> str:
        """生成学习报告"""

        self._load_latest_history()
        memory_summary = self.memory_tool.execute("summary")
        rag_stats = self.rag_tool.execute("stats")

        documents = self.history.get("documents", [])
        questions = self.history.get("questions", [])
        notes = self.history.get("notes", [])

        recent_documents = documents[-5:]
        recent_questions = questions[-10:]
        recent_notes = notes[-10:]

        doc_text = "\n".join(
            [
                f"{i + 1}. {item.get('document_name')} | {item.get('loaded_at')}"
                for i, item in enumerate(recent_documents)
            ]
        ) or "暂无导入记录"

        qa_text = "\n\n".join(
            [
                f"{i + 1}. 问题：{item.get('question')}\n回答：{str(item.get('answer'))[:300]}..."
                for i, item in enumerate(recent_questions)
            ]
        ) or "暂无问答记录"

        note_text = "\n".join(
            [
                f"{i + 1}. 【{item.get('concept') or '未命名概念'}】{item.get('note')}"
                for i, item in enumerate(recent_notes)
            ]
        ) or "暂无学习笔记"

        report = f"""
    📘 PDF 智能学习报告

    一、学习基本信息
    - user_id: {self.user_id}
    - session_id: {self.session_id}
    - 当前文档: {self.current_document or "暂无"}
    - 历史导入文档数: {len(documents)}
    - 历史提问次数: {len(questions)}
    - 历史学习笔记数: {len(notes)}
    - 历史文件路径: {self.history_path}

    二、最近导入的文档
    {doc_text}

    三、记忆系统状态
    {memory_summary}

    四、RAG 知识库状态
    {rag_stats}

    五、最近问答记录
    {qa_text}

    六、最近学习笔记
    {note_text}

    七、推荐复习方向
    1. 优先复习最近保存的学习笔记，尤其是你标注过概念名称的内容。
    2. 回顾最近问过的问题，整理其中反复出现的关键词。
    3. 对当前 PDF 继续追问：“核心概念是什么？”、“有哪些重点章节？”、“有哪些易混点？”。
    4. 对重要知识点继续添加学习笔记，形成长期语义记忆。
    """

        return report.strip()

    def _load_history(self) -> Dict[str, Any]:
        """加载学习历史"""
        return self.history_repository.load()

    def _save_history(self) -> None:
        """保存学习历史"""
        try:
            self.history["last_updated"] = datetime.now().isoformat()
            self.history_repository.save(self.history)
        except Exception as e:
            print(f"[WARNING] 学习历史保存失败: {e}")

    def _load_latest_history(self) -> Dict[str, Any]:
        with self._write_lock:
            self.history = self._load_history()
            return self.history

    def _update_history(self, mutation) -> Dict[str, Any]:
        """Merge a change into the latest persisted history under the user lock."""
        with self._write_lock:
            if self.coordinator is not None:
                self.history = self.coordinator.update_history(mutation)
            else:
                self.history = self.history_repository.update(mutation)
            return self.history

    def _execute_memory(self, action: str, **kwargs):
        with self._write_lock:
            return self.memory_tool.execute(action, **kwargs)

    def get_documents(self):
        """获取历史导入过的文档列表，用于 Gradio 下拉框"""

        documents = self._load_latest_history().get("documents", [])

        choices = []

        seen = set()

        for item in documents:
            document_id = item.get("document_id") or Path(item.get("document_path", "")).stem
            document_name = item.get("document_name") or document_id

            if not document_id:
                continue

            if document_id in seen:
                continue

            seen.add(document_id)

            label = f"{document_name} | {document_id}"
            choices.append(label)

        return choices

    def select_document(self, selected: str) -> str:
        """选择当前 PDF 文档"""

        if not selected:
            return "❌ 请选择一个 PDF 文档"

        # selected 格式：文件名 | document_id
        if "|" in selected:
            document_id = selected.split("|")[-1].strip()
        else:
            document_id = selected.strip()

        documents = self.history.get("documents", [])

        matched = None

        for item in documents:
            item_document_id = item.get("document_id") or Path(item.get("document_path", "")).stem

            if item_document_id == document_id:
                matched = item
                break

        if matched:
            self.current_document_id = document_id
            self.current_document = matched.get("document_path")

            return (
                f"✅ 当前 PDF 已切换\n"
                f"- 文档名: {matched.get('document_name')}\n"
                f"- document_id: {self.current_document_id}\n"
                f"- 路径: {self.current_document}"
            )

        self.current_document_id = document_id

        return f"✅ 当前 document_id 已设置为: {self.current_document_id}"

    def delete_current_document(self) -> str:
        """删除当前选中的 PDF：删除 RAG chunks + 删除历史文档记录 + 重置当前 PDF"""

        if not self.current_document_id:
            return "❌ 当前没有选择 PDF，无法删除"

        document_id = self.current_document_id
        result = self._delete_document_coordinated(document_id)
        self.current_document_id = None
        self.current_document = None
        return result

    def clear_all_documents(self) -> str:
        """清空全部 PDF：清空 RAG 知识库 + 清理学习历史中的文档和问答记录 + 重置当前 PDF"""

        runtime = getattr(self, "runtime", None)
        runtime_lock = getattr(runtime, "lock", None)
        with runtime_lock or self._write_lock:
            import_task_service = getattr(runtime, "import_task_service", None)
            if import_task_service is not None and import_task_service.has_active_tasks(
                self.user_id
            ):
                return "Cannot clear documents while imports are active; wait for them to finish."
            return self._clear_documents_coordinated()

    def _delete_document_coordinated(self, document_id: str) -> str:
        with self._write_lock:
            latest = self._load_history()
            source_paths = [
                Path(item.get("document_path", ""))
                for item in latest["documents"]
                if item.get("document_id") == document_id
            ]
            rag_result = self.rag_tool.execute("delete_document", document_id=document_id)
            if self.coordinator is not None:
                removed_docs, removed_questions = self.coordinator.delete_document(document_id)
                self.history = self.coordinator.load_history()
            else:
                removed_docs, removed_questions = self.history_repository.delete_document(document_id)
                self.history = self.history_repository.load()

            # Unlink source files.  When a coordinator is present every
            # path must be inside the user document root — rejections
            # are collected and reported as partial failure.
            skipped_paths: list[Path] = []
            for path in source_paths:
                if not path.exists():
                    continue
                if self.coordinator is not None:
                    try:
                        self.coordinator.safe_unlink(path)
                    except ValueError:
                        skipped_paths.append(path)
                else:
                    path.unlink()

        result = (
            f"{rag_result}\n\nHistory synchronized\n"
            f"- documents removed: {removed_docs}\n"
            f"- questions removed: {removed_questions}"
        )
        if skipped_paths:
            result += (
                "\n- ⚠️ source files outside user root were not deleted: "
                + ", ".join(str(p) for p in skipped_paths)
            )
        return result

    def _clear_documents_coordinated(self) -> str:
        with self._write_lock:
            latest = self._load_history()
            source_paths = [Path(item.get("document_path", "")) for item in latest["documents"]]
            rag_result = self.rag_tool.execute("clear")
            if self.coordinator is not None:
                removed_docs, removed_questions = self.coordinator.clear_documents()
                self.history = self.coordinator.load_history()
            else:
                removed_docs, removed_questions = self.history_repository.clear_documents()
                self.history = self.history_repository.load()

            # Unlink source files.  When a coordinator is present every
            # path must be inside the user document root — rejections
            # are collected and reported as partial failure.
            skipped_paths: list[Path] = []
            for path in source_paths:
                if not path.exists():
                    continue
                if self.coordinator is not None:
                    try:
                        self.coordinator.safe_unlink(path)
                    except ValueError:
                        skipped_paths.append(path)
                else:
                    path.unlink()

            self.current_document = None
            self.current_document_id = None
            self.stats["documents_loaded"] = 0
            self.stats["questions_asked"] = 0
        result = (
            f"{rag_result}\n\nHistory synchronized\n"
            f"- documents removed: {removed_docs}\n"
            f"- questions removed: {removed_questions}\n"
            "- notes retained"
        )
        if skipped_paths:
            result += (
                "\n- ⚠️ source files outside user root were not deleted: "
                + ", ".join(str(p) for p in skipped_paths)
            )
        return result

    def export_report_markdown(self) -> str:
        """导出学习报告为 Markdown 文件"""

        from pathlib import Path
        from datetime import datetime

        report_dir = self.report_dir
        report_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = report_dir / f"learning_report_{self.user_id}_{timestamp}.md"

        report = self.generate_report()
        if self.report_service is not None:
            with self._write_lock:
                record = self.report_service.create_markdown_snapshot(
                    self.user_id, "Learning report", report
                )
            return str(self.report_service.report_file_path(self.user_id, record.id))
        file_path.write_text(report, encoding="utf-8")
        return str(file_path)

    def export_report_docx(self, report_id: Optional[str] = None) -> str:
        """导出学习报告为格式更美观的 Word 文件"""

        from pathlib import Path
        from datetime import datetime

        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        report_dir = self.report_dir
        report_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = report_dir / f"learning_report_{self.user_id}_{timestamp}.docx"

        if self.report_service is not None:
            if report_id is None:
                report = self.generate_report()
                with self._write_lock:
                    record = self.report_service.create_markdown_snapshot(
                        self.user_id, "Learning report", report
                    )
                report_id = record.id
            else:
                report = self.report_service.read_report(self.user_id, report_id)
            file_path = self.runtime.paths.reports / f"{report_id}.docx"
            if file_path.exists():
                return str(file_path)
        else:
            report = self.generate_report()

        doc = Document()

        # =========================
        # 全局字体设置
        # =========================
        styles = doc.styles

        normal_style = styles["Normal"]
        normal_style.font.name = "微软雅黑"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        normal_style.font.size = Pt(11)

        # =========================
        # 标题
        # =========================
        title = doc.add_heading("PDF 智能学习报告", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in title.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(20)
            run.bold = True

        # =========================
        # 基本信息
        # =========================
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = info.add_run(
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}    "
            f"用户：{self.user_id}"
        )
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10)

        doc.add_paragraph("")

        # =========================
        # 正文
        # =========================
        for raw_line in report.splitlines():
            line = raw_line.strip()

            if not line:
                doc.add_paragraph("")
                continue

            # 去掉开头的报告标题，避免重复
            if line.startswith("📘 PDF 智能学习报告"):
                continue

            # 一级小节标题
            if (
                    line.startswith("一、")
                    or line.startswith("二、")
                    or line.startswith("三、")
                    or line.startswith("四、")
                    or line.startswith("五、")
                    or line.startswith("六、")
                    or line.startswith("七、")
                    or line.startswith("八、")
                    or line.startswith("九、")
                    or line.startswith("十、")
            ):
                heading = doc.add_heading(line, level=2)

                for run in heading.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(15)
                    run.bold = True

                continue

            # 项目符号
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(line[2:])
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(11)
                continue

            # 数字列表
            if len(line) > 2 and line[0].isdigit() and line[1] in [".", "．", "、"]:
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(line)
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(11)
                continue

            # 普通段落
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25

            run = p.add_run(line)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(11)

        # =========================
        # 页脚式结尾
        # =========================
        doc.add_paragraph("")
        end = doc.add_paragraph("—— 由 PDF 智能学习助手自动生成 ——")
        end.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in end.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(9)
            run.italic = True

        doc.save(str(file_path))

        return str(file_path)
